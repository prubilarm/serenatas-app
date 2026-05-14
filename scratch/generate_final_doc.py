from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell borders.
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#000000"}, ...)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ('top', 'start', 'bottom', 'end'):
        if side in kwargs:
            tag = 'w:{}'.format(side)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for attr, val in kwargs[side].items():
                element.set(qn('w:{}'.format(attr)), str(val))

def create_final_document():
    doc = Document()

    # Configuración de márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- ENCABEZADO ---
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_h = header.add_run('Actividad de Abstracción Completa\nSistema: "El Mariachi Aventurero" v4.0')
    run_h.font.size = Pt(10)
    run_h.font.color.rgb = RGBColor(128, 128, 128)

    # ==========================================
    # PARTE A: CARA FRONTAL
    # ==========================================
    title_a = doc.add_heading('A. CARA FRONTAL: HISTORIA Y PROCESOS', 0)
    title_a.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1. Contexto del Proyecto', level=1)
    context = doc.add_paragraph()
    context.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    context.add_run('El sistema ')
    context.add_run('"El Mariachi Aventurero"').bold = True

    context.add_run(' es una solución de gestión interna diseñada para resolver la desorganización en la administración de servicios musicales. Este proyecto está concebido como una herramienta de uso exclusivo para el ')
    context.add_run('músico').underline = True
    context.add_run(' (administrador), permitiéndole llevar un registro histórico de cada ')
    context.add_run('cliente').underline = True
    context.add_run(', coordinar la logística de las ')
    context.add_run('serenatas').underline = True
    context.add_run(' y asegurar el flujo de caja mediante el control de ')
    context.add_run('pagos').underline = True
    context.add_run(' y ')
    context.add_run('comprobantes').underline = True
    context.add_run('.')

    doc.add_heading('2. Procesos del Sistema (Narrativa)', level=1)
    procs = [
        ("Registro de Cliente", "Se capturan los datos de contacto iniciales para crear un perfil único."),
        ("Agendamiento", "Se vincula una serenata al cliente con datos de fecha, hora, dirección y festejada."),
        ("Ejecución y Estado", "Actualización del progreso del servicio (En camino, Realizada)."),
        ("Gestión Financiera", "Registro de transacciones, métodos de pago y validación de saldos.")
    ]
    for t, d in procs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{t}: ').bold = True
        p.add_run(d)

    doc.add_page_break()

    # ==========================================
    # PARTE B: CARA POSTERIOR
    # ==========================================
    title_b = doc.add_heading('B. CARA POSTERIOR: DIAGRAMA MER', 0)
    title_b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1. Diagramación (Modelo Relacional Lógico)', level=1)
    doc.add_paragraph('A continuación se representan las entidades, sus atributos y las relaciones que sostienen el sistema:')

    # --- TABLA CLIENTE ---
    doc.add_heading('Entidad: CLIENTE', level=2)
    table_c = doc.add_table(rows=5, cols=3)
    table_c.style = 'Table Grid'
    hdr_cells = table_c.rows[0].cells
    hdr_cells[0].text = 'Campo'
    hdr_cells[1].text = 'Tipo'
    hdr_cells[2].text = 'Key'
    for cell in hdr_cells: cell.paragraphs[0].runs[0].bold = True

    data_c = [
        ("id", "UUID", "PK"),
        ("nombre", "TEXT", ""),
        ("telefono", "TEXT", ""),
        ("observaciones", "TEXT", "")
    ]
    for i, (f, t, k) in enumerate(data_c):
        row = table_c.rows[i+1].cells
        row[0].text, row[1].text, row[2].text = f, t, k

    doc.add_paragraph() # Spacer

    # --- TABLA SERENATA ---
    doc.add_heading('Entidad: SERENATA', level=2)
    table_s = doc.add_table(rows=9, cols=3)
    table_s.style = 'Table Grid'
    hdr_s = table_s.rows[0].cells
    hdr_s[0].text, hdr_s[1].text, hdr_s[2].text = 'Campo', 'Tipo', 'Key'
    for cell in hdr_s: cell.paragraphs[0].runs[0].bold = True

    data_s = [
        ("id", "UUID", "PK"),
        ("cliente_id", "UUID", "FK (CLIENTE)"),
        ("nombre_festejada", "TEXT", ""),
        ("fecha", "DATE", ""),
        ("hora", "TIME", ""),
        ("comuna", "TEXT", ""),
        ("precio_total", "NUMERIC", ""),
        ("estado", "TEXT", "")
    ]
    for i, (f, t, k) in enumerate(data_s):
        row = table_s.rows[i+1].cells
        row[0].text, row[1].text, row[2].text = f, t, k

    doc.add_paragraph()

    # --- TABLA PAGO ---
    doc.add_heading('Entidad: PAGO', level=2)
    table_p = doc.add_table(rows=5, cols=3)
    table_p.style = 'Table Grid'
    hdr_p = table_p.rows[0].cells
    hdr_p[0].text, hdr_p[1].text, hdr_p[2].text = 'Campo', 'Tipo', 'Key'
    for cell in hdr_p: cell.paragraphs[0].runs[0].bold = True

    data_p = [
        ("id", "UUID", "PK"),
        ("serenata_id", "UUID", "FK (SERENATA)"),
        ("monto", "NUMERIC", ""),
        ("metodo", "TEXT", "")
    ]
    for i, (f, t, k) in enumerate(data_p):
        row = table_p.rows[i+1].cells
        row[0].text, row[1].text, row[2].text = f, t, k

    # --- CARDINALIDADES ---
    doc.add_heading('2. Cardinalidades y Relaciones', level=1)
    rel = doc.add_paragraph()
    rel.add_run('• Relación CLIENTE - SERENATA (1:N): ').bold = True
    rel.add_run('Un cliente puede registrar múltiples serenatas en el sistema, pero cada serenata está asociada obligatoriamente a un solo cliente para fines de facturación y contacto.')
    
    rel2 = doc.add_paragraph()
    rel2.add_run('• Relación SERENATA - PAGO (1:N): ').bold = True
    rel2.add_run('Una serenata puede recibir varios pagos (ej. reserva y saldo final), permitiendo un control de abonos dinámico.')

    # Footer
    doc.add_paragraph('\n' * 1)
    footer = doc.add_paragraph('Documentación técnica generada para el Proyecto Académico "Mariachi Aventurero".')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)


    # Save
    filename = 'Actividad_Abstraccion_Serenatas_FINAL.docx'
    doc.save(filename)
    print(f'Documento final guardado como: {filename}')

if __name__ == '__main__':
    create_final_document()
