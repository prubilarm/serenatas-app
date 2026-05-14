from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_chen_style_document():
    doc = Document()

    # Configuración de márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # --- ENCABEZADO ---
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_h = header.add_run('Actividad de Abstracción: El Mariachi Aventurero\nDiagrama Modelo Entidad-Relación (MER)')
    run_h.font.size = Pt(9)
    run_h.font.color.rgb = RGBColor(100, 100, 100)

    # ==========================================
    # PARTE A: CARA FRONTAL
    # ==========================================
    doc.add_heading('A. CARA FRONTAL: HISTORIA Y PROCESOS', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('1. Contexto del Proyecto', level=1)
    p_ctx = doc.add_paragraph()
    p_ctx.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_ctx.add_run('El proyecto ')
    p_ctx.add_run('"El Mariachi Aventurero"').bold = True

    p_ctx.add_run(' es un sistema de gestión privada diseñado para optimizar el orden administrativo del músico. La propuesta se enfoca en centralizar la información de ')
    p_ctx.add_run('clientes').underline = True
    p_ctx.add_run(' y ')
    p_ctx.add_run('serenatas').underline = True
    p_ctx.add_run(' para evitar traslapes de horarios y asegurar el cobro de cada servicio mediante un registro riguroso de ')
    p_ctx.add_run('pagos').underline = True
    p_ctx.add_run('. Es una herramienta personal, no orientada al cliente final, sino a la eficiencia del propietario.')

    doc.add_heading('2. Procesos del Sistema (Narrativa de Datos)', level=1)
    p_proc = doc.add_paragraph()
    p_proc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_proc.add_run('El flujo inicia con el registro del ')
    p_proc.add_run('cliente').underline = True
    p_proc.add_run(', capturando su nombre y teléfono. Luego, se agenda la ')
    p_proc.add_run('serenata').underline = True
    p_proc.add_run(' indicando fecha, hora, dirección y la festejada. El ')
    p_proc.add_run('músico').underline = True
    p_proc.add_run(' gestiona el estado del evento (consulta, confirmada, realizada) y finalmente registra el ')
    p_proc.add_run('pago').underline = True
    p_proc.add_run(' recibido, validando el ')
    p_proc.add_run('monto').underline = True
    p_proc.add_run(' y el ')
    p_proc.add_run('comprobante').underline = True
    p_proc.add_run('.')

    doc.add_page_break()

    # ==========================================
    # PARTE B: CARA POSTERIOR
    # ==========================================
    doc.add_heading('B. CARA POSTERIOR: DIAGRAMA MER', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1. Diagramación (Notación de Chen)', level=1)
    doc.add_paragraph('Representación gráfica de entidades (rectángulos) y relaciones (rombos):').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insertar la imagen del MER estilo Chen
    image_path = r'C:\Users\prubi\.gemini\antigravity\brain\7387ae71-aba8-4bf4-b378-d6579c5aeff8\diagrama_chen_serenatas_1778780356099.png'
    try:
        doc.add_picture(image_path, width=Inches(6.2))
    except Exception:
        doc.add_paragraph("[Diagrama MER no disponible automáticamente - Por favor inserta la imagen generada aquí]").italic = True

    # --- LEYENDA ---
    doc.add_paragraph('\n')
    legend_table = doc.add_table(rows=1, cols=1)
    legend_table.style = 'Table Grid'
    cell = legend_table.rows[0].cells[0]
    p_leg = cell.paragraphs[0]
    p_leg.add_run('LEYENDA DEL MODELO').bold = True
    cell.add_paragraph('■ Rectángulo: Entidad (Objetos del sistema)')
    cell.add_paragraph('◆ Rombo: Relación (Acciones o vínculos)')
    cell.add_paragraph('— Línea: Conexión entre entidades y relaciones')
    cell.add_paragraph('1 : N: Cardinalidad (Uno a Muchos)')

    # --- NOTAS ---
    doc.add_heading('2. Reglas de Negocio y Cardinalidad', level=1)
    notes = [
        "Un Cliente puede solicitar muchas (N) Serenatas.",
        "Una Serenata es gestionada por un único (1) Músico administrador.",
        "Una Serenata genera uno o varios (N) Pagos asociados.",
        "Cada relación define un flujo de datos lógico en la base de datos."
    ]
    for note in notes:
        doc.add_paragraph(note, style='List Bullet')

    # Guardar
    final_filename = 'Trabajo_Serenatas_MER_Chen.docx'
    doc.save(final_filename)
    print(f'Documento final guardado como: {final_filename}')

if __name__ == '__main__':
    create_chen_style_document()
