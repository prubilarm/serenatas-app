from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_premium_document():
    doc = Document()

    # Configuración de márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- ENCABEZADO ---
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_h = header.add_run('Actividad de Abstracción: Narrativa y Procesos\nProyecto: Sistema de Gestión "El Mariachi Aventurero"')
    run_h.font.size = Pt(10)
    run_h.font.color.rgb = RGBColor(128, 128, 128)

    # --- TÍTULO PRINCIPAL ---
    title = doc.add_heading('CARA FRONTAL: HISTORIA Y PROCESOS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- CONTEXTO DEL PROYECTO ---
    doc.add_heading('1. Contexto del Proyecto', level=1)
    
    context = doc.add_paragraph()
    context.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    context.add_run('El presente proyecto nace de la necesidad crítica de profesionalizar y centralizar la gestión de servicios musicales de ')
    context.add_run('"El Mariachi Aventurero"').bold = True
    context.add_run('. A menudo, los músicos independientes enfrentan el caos de manejar múltiples ')
    context.add_run('serenatas').underline = True
    context.add_run(' a través de canales informales, lo que conlleva a errores en horarios, direcciones o cobros.\n\n')
    
    context.add_run('La propuesta consiste en una plataforma de administración personalizada (Back-office) diseñada específicamente para el ')
    context.add_run('músico').underline = True
    context.add_run(' (dueño del negocio). No se trata de una aplicación para que el público reserve, sino de un centro de control donde el administrador puede registrar cada interacción con el ')
    context.add_run('cliente').underline = True
    context.add_run(', programar la logística de los eventos y realizar un seguimiento financiero en tiempo real. El objetivo es sustituir las notas físicas por un ecosistema digital que garantice que ninguna ')
    context.add_run('presentación').underline = True
    context.add_run(' sea olvidada y que cada ')
    context.add_run('pago').underline = True
    context.add_run(' esté debidamente respaldado por un ')
    context.add_run('comprobante').underline = True
    context.add_run(' digital.')

    # --- PROCESOS DEL SISTEMA ---
    doc.add_heading('2. Procesos del Sistema (Flujo de Datos)', level=1)
    
    processes = [
        ("Captura y Registro de Interés", 
         "El flujo inicia cuando el músico recibe un contacto. Se procede a registrar al cliente en la base de datos, almacenando su nombre y teléfono. Este paso genera un historial que permite identificar clientes frecuentes y sus preferencias anteriores."),
        
        ("Configuración Logística de la Serenata", 
         "Una vez pactado el servicio, se crea un registro de serenata vinculado al cliente. Se capturan datos dinámicos: nombre de la festejada, motivo (cumpleaños, aniversario, etc.), fecha exacta, hora de llegada y la dirección precisa con su respectiva comuna para optimizar el traslado."),
        
        ("Operación y Control de Estados", 
         "El sistema permite la trazabilidad del servicio mediante estados. La serenata inicia como 'consulta', pasa a 'confirmada' y durante el día del evento el músico actualiza a 'en camino' y 'realizada'. Esto alimenta un calendario operativo que evita traslapes de horarios."),
        
        ("Gestión de Transacciones Financieras", 
         "Finalizada la serenata (o en el momento de la reserva), el flujo de datos se dirige al módulo de pagos. Se registra el monto recibido, el método utilizado (efectivo o transferencia) y se carga la URL de la imagen del comprobante si aplica. Esto asegura que la contabilidad del negocio sea transparente."),
        
        ("Análisis de Rendimiento y Reportes", 
         "Los datos de serenatas y pagos se consolidan en reportes periódicos. El sistema permite al músico visualizar sus ingresos totales, la efectividad de sus cobros y las comunas de mayor demanda para ajustar su estrategia de marketing y logística.")
    ]

    for p_title, p_desc in processes:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{p_title}: ').bold = True
        p.add_run(p_desc)

    # --- IDENTIFICACIÓN DE ENTIDADES ---
    doc.add_heading('3. Identificación de Entidades', level=1)
    
    ent_desc = doc.add_paragraph()
    ent_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ent_desc.add_run('Para la construcción del modelo de datos, se han extraído los siguientes sustantivos clave de la narrativa anterior, los cuales representan las entidades fundamentales del sistema:')

    # Lista de entidades con descripción breve
    entity_list = [
        ("Cliente", "Persona que solicita el servicio y mantiene el vínculo comercial."),
        ("Serenata", "El evento o servicio principal que se agenda y ejecuta."),
        ("Pago", "La transacción financiera vinculada a una o varias serenatas."),
        ("Músico", "El usuario administrador que gestiona el sistema."),
        ("Comprobante", "Documento digital o imagen que respalda un pago."),
        ("Festejada", "Persona que recibe el homenaje musical (dato relevante para la personalización)."),
        ("Comuna", "Entidad geográfica que permite clasificar y filtrar la logística de servicios.")
    ]

    for e_name, e_info in entity_list:
        p_e = doc.add_paragraph()
        run_e = p_e.add_run(f"• {e_name}")
        run_e.bold = True
        run_e.underline = True
        p_e.add_run(f": {e_info}")

    # --- PIE DE PÁGINA ---
    doc.add_paragraph('\n' * 2)
    note = doc.add_paragraph('Nota: Este análisis servirá de base para el diseño del Diagrama Modelo Entidad-Relación (MER) en la siguiente etapa.')
    note.italic = True
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    filename = 'Actividad_Serenatas_Parte_A_Completa.docx'
    doc.save(filename)
    print(f'Documento detallado guardado como: {filename}')

if __name__ == '__main__':
    create_premium_document()
