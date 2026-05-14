from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()

    # Title
    title = doc.add_heading('Actividad de Aula: Sistema de Gestión de Serenatas', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    doc.add_paragraph('Este documento presenta el análisis y la estructura del proyecto "Mariachi Aventurero", enfocado en la organización interna de servicios musicales.')

    # Section A
    doc.add_heading('A. CARA FRONTAL: HISTORIA Y PROCESOS', level=1)

    # 1. Contexto del Proyecto
    doc.add_heading('1. Contexto del Proyecto', level=2)
    p1 = doc.add_paragraph()
    p1.add_run('El proyecto ').bold = True
    p1.add_run('"Serenatas App" (Mariachi Aventurero) ')
    p1.add_run('es una plataforma de gestión personalizada diseñada para optimizar la organización interna de las presentaciones musicales. A diferencia de una aplicación orientada al cliente final, esta herramienta está enfocada exclusivamente en el uso administrativo del músico (el propietario), permitiéndole llevar un control riguroso de su agenda, clientes y finanzas. La propuesta busca centralizar la información que antes se manejaba de forma dispersa, asegurando que cada ')
    p1.add_run('serenata').underline = True
    p1.add_run(' esté debidamente documentada, desde el contacto inicial hasta el ')
    p1.add_run('pago').underline = True
    p1.add_run(' final.')

    # 2. Procesos del Sistema
    doc.add_heading('2. Procesos del Sistema (Narrativa de flujo de datos)', level=2)
    
    proc_text = [
        ('Registro de Cliente', 'El proceso inicia cuando el músico recibe una solicitud. Se registra al cliente con su nombre y teléfono para mantener una base de datos de contactos histórica.'),
        ('Agendamiento de Serenata', 'Se crea una nueva serenata vinculada al cliente. Aquí se capturan detalles críticos: nombre de la festejada, motivo, fecha, hora, dirección y comuna. El sistema permite clasificar el tipo (express o full) y asignar un precio total.'),
        ('Seguimiento de Estado', 'La serenata transita por diversos estados (consulta, cotizada, confirmada, en camino, realizada, pagada, cancelada), lo que permite al músico saber exactamente qué eventos están pendientes, cuáles requieren atención inmediata o cuáles están listos para ser cobrados.'),
        ('Gestión de Pagos', 'Una vez realizada la presentación, se registra el pago. Se especifica el monto, el método (efectivo o transferencia) y se puede adjuntar un comprobante digital (URL de imagen).'),
        ('Generación de Reportes', 'El sistema procesa los datos acumulados para mostrar resúmenes de ingresos mensuales, cantidad de serenatas por tipo y efectividad de cobros.')
    ]

    for title, desc in proc_text:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)

    # 3. Identificación de Entidades
    doc.add_heading('3. Identificación de Entidades (Sustantivos clave)', level=2)
    doc.add_paragraph('A continuación se listan las entidades principales identificadas en el flujo del sistema, las cuales formarán la base del modelo relacional:')
    
    entities = ['Cliente', 'Serenata', 'Pago', 'Festejada', 'Músico', 'Comprobante', 'Estado']
    p_entities = doc.add_paragraph()
    for i, entity in enumerate(entities):
        run = p_entities.add_run(entity)
        run.underline = True
        if i < len(entities) - 1:
            p_entities.add_run(', ')

    # Save
    filename = 'Actividad_Serenatas_Parte_A.docx'
    doc.save(filename)
    print(f'Documento guardado como: {filename}')

if __name__ == '__main__':
    create_document()
