from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_styled_document():
    doc = Document()

    # Márgenes tipo oficio/profesional
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # --- PARTE A: CARA FRONTAL ---
    doc.add_heading('A. CARA FRONTAL: HISTORIA Y PROCESOS', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('1. Contexto del Proyecto', level=1)
    p_ctx = doc.add_paragraph()
    p_ctx.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_ctx.add_run('El sistema ')
    p_ctx.add_run('"El Mariachi Aventurero"').bold = True

    p_ctx.add_run(' es una plataforma de gestión personalizada diseñada para la administración interna de servicios musicales. A diferencia de apps comerciales, esta herramienta es de uso exclusivo para el ')
    p_ctx.add_run('músico').underline = True
    p_ctx.add_run(' (propietario), permitiéndole organizar su agenda de ')
    p_ctx.add_run('serenatas').underline = True
    p_ctx.add_run(', gestionar datos de ')
    p_ctx.add_run('clientes').underline = True
    p_ctx.add_run(' y llevar un control financiero mediante el registro de ')
    p_ctx.add_run('pagos').underline = True
    p_ctx.add_run(' y ')
    p_ctx.add_run('comprobantes').underline = True
    p_ctx.add_run('.')

    doc.add_heading('2. Procesos del Sistema (Narrativa Detallada)', level=1)
    procs = [
        ("Identificación de Demanda", "El músico recibe una solicitud y registra al cliente (Nombre, Teléfono) para iniciar el historial."),
        ("Planificación Técnica", "Se agenda la serenata vinculándola al cliente. Se definen datos críticos como fecha, hora, dirección y la festejada."),
        ("Control Logístico", "Se actualiza el estado (Consulta -> Confirmada -> Realizada) para coordinar el traslado y ejecución."),
        ("Cierre Financiero", "Se registra el pago del servicio, permitiendo adjuntar comprobantes y liquidar saldos pendientes.")
    ]
    for t, d in procs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{t}: ').bold = True
        p.add_run(d)

    doc.add_page_break()

    # --- PARTE B: CARA POSTERIOR ---
    doc.add_heading('B. CARA POSTERIOR: DIAGRAMA MER', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1. Diagramación (Estilo Chen / Relacional)', level=1)
    
    # Insertar la imagen generada (usando la ruta absoluta proporcionada)
    image_path = r'C:\Users\prubi\.gemini\antigravity\brain\7387ae71-aba8-4bf4-b378-d6579c5aeff8\diagrama_er_serenatas_1778780184410.png'
    try:
        doc.add_picture(image_path, width=Inches(6.5))
    except Exception as e:
        doc.add_paragraph(f"[Espacio para Diagrama MER: {image_path}]")
        doc.add_paragraph("(Error al insertar imagen automáticamente, por favor inserta el archivo generado manualmente)").italic = True

    # --- LEYENDA (Al estilo de la foto) ---
    doc.add_paragraph()
    table_l = doc.add_table(rows=1, cols=1)
    table_l.style = 'Table Grid'
    cell_l = table_l.rows[0].cells[0]
    p_l = cell_l.paragraphs[0]
    p_l.add_run('LEYENDA').bold = True
    cell_l.add_paragraph('⬚  Entidad (Tabla)')
    cell_l.add_paragraph('◇  Relación (Acción)')
    cell_l.add_paragraph('PK  Clave Primaria')
    cell_l.add_paragraph('FK  Clave Foránea')
    cell_l.add_paragraph('1 / N  Cardinalidad (Uno a Muchos)')

    # --- NOTAS DE NEGOCIO ---
    doc.add_heading('2. Notas y Reglas de Cardinalidad', level=1)
    notes_box = doc.add_table(rows=1, cols=1)
    notes_box.style = 'Table Grid'
    cell_n = notes_box.rows[0].cells[0]
    cell_n.add_paragraph('• Un Cliente puede solicitar muchas (N) Serenatas en el tiempo.')
    cell_n.add_paragraph('• Una Serenata pertenece a un único (1) Cliente.')
    cell_n.add_paragraph('• Una Serenata puede generar varios (N) Pagos (abonos o totales).')
    cell_n.add_paragraph('• Un Músico (Administrador) gestiona todos los procesos del sistema.')

    # Footer
    footer = doc.add_paragraph('\nProyecto: El Mariachi Aventurero - Documentación Académica')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)

    # Save
    filename = 'Actividad_Serenatas_ESTILO_FOTO.docx'
    doc.save(filename)
    print(f'Documento final estilo foto guardado como: {filename}')

if __name__ == '__main__':
    create_styled_document()
