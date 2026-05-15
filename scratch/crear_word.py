from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_word():
    doc = Document()

    # Título Principal
    title = doc.add_heading('Informe de Gestión: El Mariachi Aventurero', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1. Contexto del Proyecto', level=1)
    doc.add_paragraph(
        'El proyecto “El Mariachi Aventurero” es una aplicación de gestión interna creada para administrar servicios de serenatas '
        'ofrecidos por un músico independiente en la ciudad de Los Ángeles, Chile. '
        'La finalidad del sistema es organizar en una sola plataforma la información de los clientes, los datos de cada serenata, '
        'la agenda de eventos y el registro de pagos.'
    )

    doc.add_heading('2. Procesos del Sistema', level=1)
    doc.add_paragraph(
        'El funcionamiento del sistema comienza cuando el músico recibe una solicitud. Primero, se registran los datos del cliente '
        '(nombre, teléfono y observaciones). Luego, se crea una serenata asociada a ese cliente.'
    )
    doc.add_paragraph(
        'En el registro de la serenata se ingresan los datos principales: nombre de la festejada, motivo, fecha, hora, dirección, '
        'comuna, mensaje especial, tipo, canciones y precio total. Una vez finalizada, se registra el pago correspondiente '
        '(monto, método, comprobante).'
    )

    doc.add_heading('3. Identificación de Entidades', level=1)
    doc.add_heading('Entidad CLIENTES', level=2)
    doc.add_paragraph('Atributos: id, nombre, telefono, observaciones, created_at')

    doc.add_heading('Entidad SERENATAS', level=2)
    doc.add_paragraph('Atributos: id, cliente_id, nombre_festejada, motivo, fecha, hora, direccion, comuna, tipo, precio_total, estado.')

    doc.add_heading('Entidad PAGOS', level=2)
    doc.add_paragraph('Atributos: id, serenata_id, monto, metodo, comprobante_url, fecha_pago.')

    doc.add_heading('4. Relaciones Identificadas', level=1)
    doc.add_paragraph('• CLIENTES 1:N SERENATAS: Un cliente puede solicitar varias serenatas.')
    doc.add_paragraph('• SERENATAS 1:N PAGOS: Una serenata puede tener varios pagos (abonos).')

    doc.add_heading('5. Conclusión', level=1)
    doc.add_paragraph(
        'En conclusión, el sistema se organiza alrededor de la entidad serenatas, ya que desde ella se conectan los datos del '
        'cliente y los pagos realizados. Las pantallas de agenda y calendario son vistas de la información almacenada.'
    )

    doc.save('Documentacion_Mariachi_Final.docx')
    print('Archivo Documentacion_Mariachi_Final.docx generado con éxito.')

if __name__ == '__main__':
    generate_word()
