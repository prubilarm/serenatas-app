from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

IMG_PATH = r'C:\Users\prubi\.gemini\antigravity\brain\7387ae71-aba8-4bf4-b378-d6579c5aeff8\mer_serenatas_v2_completo_1778781120319.png'
OUTPUT = 'Trabajo_Final_Serenatas_MER_v2.docx'

def heading(doc, text, level):
    h = doc.add_heading(text, level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def justified(doc, *runs):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for item in runs:
        if isinstance(item, str):
            p.add_run(item)
        elif isinstance(item, dict):
            r = p.add_run(item['text'])
            r.bold = item.get('bold', False)
            r.underline = item.get('underline', False)
            r.italic = item.get('italic', False)
    return p

def bullet(doc, label, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f'{label}: ')
    r.bold = True
    p.add_run(text)
    return p

def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ENCABEZADO
    ph = doc.add_paragraph()
    ph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rh = ph.add_run('Actividad de Aula - Abstraccion de Datos\nProyecto: El Mariachi Aventurero v4.0')
    rh.font.size = Pt(9)
    rh.font.color.rgb = RGBColor(100, 100, 100)

    # ==========================================
    # CARA A
    # ==========================================
    ta = doc.add_heading('A     CARA FRONTAL: HISTORIA Y PROCESOS', 0)
    ta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading(doc, '1. Contexto del Proyecto', 1)
    justified(doc,
        'El proyecto ',
        {'text': '"El Mariachi Aventurero"', 'bold': True},
        ' es una solucion tecnologica de gestion interna disenada para el musico-empresario. '
        'A diferencia de plataformas de reserva para el publico, esta herramienta esta orientada exclusivamente '
        'al ',
        {'text': 'administrador', 'underline': True},
        ' del negocio, es decir, ',
        {'text': 'el propio musico', 'bold': True},
        '. Su objetivo es eliminar el caos de gestionar ',
        {'text': 'serenatas', 'underline': True},
        ' por WhatsApp o anotaciones fisicas, centralizando en una sola plataforma digital los datos de sus ',
        {'text': 'clientes', 'underline': True},
        ', los detalles logisticos de cada presentacion y el flujo de ',
        {'text': 'pagos', 'underline': True},
        ' con sus respectivos ',
        {'text': 'comprobantes', 'underline': True},
        '.'
    )

    justified(doc,
        'La plataforma cuenta con tres interfaces: una ',
        {'text': 'aplicacion web (panel administrativo)', 'bold': True},
        ', una ',
        {'text': 'aplicacion movil', 'bold': True},
        ' para el musico en terreno, y una ',
        {'text': 'API REST', 'bold': True},
        ' que conecta ambas con la base de datos en Supabase. El ',
        {'text': 'musico', 'underline': True},
        ' puede verificar su agenda desde el celular mientras se traslada entre ',
        {'text': 'serenatas', 'underline': True},
        ' y consultar reportes de ingresos desde su computador.'
    )

    heading(doc, '2. Procesos del Sistema (Narrativa del Flujo de Datos)', 1)

    justified(doc,
        'El sistema opera bajo un flujo secuencial que cubre el ciclo de vida completo de cada ',
        {'text': 'serenata', 'underline': True},
        '. Cada vez que el ',
        {'text': 'musico', 'underline': True},
        ' recibe una solicitud, se desencadena una cadena de procesos:'
    )

    bullet(doc, 'Registro del Cliente',
        'El sistema verifica si el contacto ya existe por telefono; si no, crea un nuevo perfil '
        'de cliente con nombre, telefono y observaciones opcionales.')

    bullet(doc, 'Eleccion del Tipo de Serenata',
        'El musico define el tipo de servicio: '
        'Express ($25.000, 2 canciones a eleccion), '
        'Full ($40.000, 4 canciones a eleccion), o '
        'Personalizado (precio segun distancia si es fuera de Los Angeles, canciones a convenir). '
        'El sistema registra el tipo, las canciones elegidas, el valor base y si aplica recargo por distancia.')

    bullet(doc, 'Agendamiento de la Serenata',
        'Se registra: nombre de la festejada, motivo, fecha, hora exacta, direccion, '
        'comuna y un mensaje especial opcional.')

    bullet(doc, 'Seguimiento por Estados',
        'La serenata avanza por: consulta -> cotizada -> confirmada -> "en camino" -> realizada -> pagada / cancelada. '
        'Esto permite coordinar la logistica y evitar traslapes de horarios.')

    bullet(doc, 'Registro de Transacciones',
        'Al concretar el servicio se registra el pago con monto, metodo (efectivo o transferencia) '
        'y se adjunta la URL del comprobante digital.')

    bullet(doc, 'Generacion de Reportes',
        'El sistema consolida pagos y serenatas para mostrar ingresos por periodo, '
        'tipos de servicio mas solicitados y comunas con mayor demanda.')

    heading(doc, '3. Identificacion de Entidades (Sustantivos Clave)', 1)

    entities = [
        ('Cliente',     'Persona que contacta y solicita el servicio musical.'),
        ('Serenata',    'El evento musical central con tipo, canciones y precio definidos.'),
        ('Pago',        'Transaccion economica que cierra el ciclo de servicio.'),
        ('Musico',      'El administrador del sistema (propietario del negocio).'),
        ('Festejada',   'Persona que recibe el homenaje musical.'),
        ('Comprobante', 'Documento digital que respalda un pago registrado.'),
        ('Estado',      'Etapa actual del servicio (confirmada, realizada, pagada...).'),
        ('Cancion',     'Pieza musical elegida segun el tipo de serenata.'),
    ]
    for name, desc in entities:
        p = doc.add_paragraph()
        rn = p.add_run(f'* {name}')
        rn.bold = True
        rn.underline = True
        p.add_run(f': {desc}')

    # ==========================================
    # CARA B
    # ==========================================
    doc.add_page_break()

    tb = doc.add_heading('B     CARA POSTERIOR: DIAGRAMA MER', 0)
    tb.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading(doc, '1. Diagramacion - Notacion de Chen', 1)
    justified(doc,
        'El diagrama representa el Modelo Entidad-Relacion (MER) completo, usando la notacion de Chen: '
        'rectangulos para entidades, ovalos para atributos y rombos para las relaciones. '
        'Los atributos ',
        {'text': 'PK', 'bold': True},
        ' (subrayados) son llaves primarias y los ',
        {'text': 'FK', 'bold': True},
        ' son llaves foraneas que vinculan las tablas.'
    )

    doc.add_paragraph()

    try:
        doc.add_picture(IMG_PATH, width=Inches(6.2))
        cap = doc.add_paragraph('Diagrama M.E.R. - Sistema El Mariachi Aventurero v4.0')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].italic = True
    except Exception as e:
        doc.add_paragraph(f'[ERROR insertando imagen: {e}]')

    doc.add_paragraph()

    heading(doc, '2. Atributos y Llaves Primarias / Foraneas', 1)

    attr_data = [
        ('CLIENTE', [
            ('id',           'UUID',                   'PK'),
            ('nombre',       'TEXT',                   ''),
            ('telefono',     'TEXT',                   ''),
            ('observaciones','TEXT',                   ''),
            ('created_at',   'TIMESTAMP',              ''),
        ]),
        ('SERENATA', [
            ('id',              'UUID',                          'PK'),
            ('cliente_id',      'UUID',                          'FK -> CLIENTE'),
            ('nombre_festejada','TEXT',                           ''),
            ('motivo',          'TEXT',                           ''),
            ('fecha',           'DATE',                           ''),
            ('hora',            'TIME',                           ''),
            ('direccion',       'TEXT',                           ''),
            ('comuna',          'TEXT',                           ''),
            ('tipo',            "TEXT ('express'|'full'|'personalizado')", ''),
            ('canciones',       'INTEGER  (express=2, full=4, personalizado=variable)', ''),
            ('valor_base',      'NUMERIC  (express=$25.000, full=$40.000, personalizado=segun distancia)', ''),
            ('fuera_ciudad',    'BOOLEAN  (TRUE si es fuera de Los Angeles => recargo)', ''),
            ('precio_total',    'NUMERIC',                        ''),
            ('mensaje_especial','TEXT',                           ''),
            ('estado',          "TEXT ('consulta'...'pagada')",   ''),
            ('created_at',      'TIMESTAMP',                     ''),
        ]),
        ('PAGO', [
            ('id',             'UUID',                                    'PK'),
            ('serenata_id',    'UUID',                                    'FK -> SERENATA'),
            ('monto',          'NUMERIC',                                  ''),
            ('metodo',         "TEXT ('efectivo'|'transferencia')",        ''),
            ('comprobante_url','TEXT',                                     ''),
            ('fecha_pago',     'TIMESTAMP',                               ''),
        ]),
    ]

    for entity_name, fields in attr_data:
        doc.add_paragraph()
        ph_e = doc.add_paragraph()
        ph_e.add_run(f'Entidad: {entity_name}').bold = True

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = 'Atributo', 'Tipo de Dato / Valores', 'Llave'
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True

        for fname, ftype, fkey in fields:
            row = table.add_row().cells
            r0 = row[0].paragraphs[0].add_run(fname)
            if fkey.startswith('PK'):
                r0.underline = True
                r0.bold = True
            row[1].text = ftype
            row[2].text = fkey

    doc.add_paragraph()
    heading(doc, '3. Reglas de Negocio y Cardinalidades', 1)

    cards = [
        ('CLIENTE - SERENATA  (1:N)',
         'Un cliente puede solicitar multiples serenatas. Cada serenata pertenece a un unico cliente (SET NULL si se borra el cliente).'),
        ('SERENATA - PAGO  (1:N)',
         'Una serenata puede tener varios pagos (abonos). Si la serenata es eliminada, sus pagos se eliminan en cascada (ON DELETE CASCADE).'),
    ]
    for tc, dc in cards:
        pc = doc.add_paragraph()
        pc.add_run(f'* {tc}: ').bold = True
        pc.add_run(dc)

    doc.add_paragraph()
    heading(doc, '4. Reglas de Precio por Tipo de Serenata', 1)

    price_table = doc.add_table(rows=4, cols=4)
    price_table.style = 'Table Grid'
    headers = price_table.rows[0].cells
    headers[0].text = 'Tipo'
    headers[1].text = 'Valor Base'
    headers[2].text = 'Canciones'
    headers[3].text = 'Condicion'
    for cell in headers:
        cell.paragraphs[0].runs[0].bold = True

    rows_data = [
        ('Express',      '$25.000', '2 (a eleccion)', 'Dentro de Los Angeles'),
        ('Full',         '$40.000', '4 (a eleccion)', 'Dentro de Los Angeles'),
        ('Personalizado','Variable','A convenir',     'Fuera de Los Angeles (recargo por distancia)'),
    ]
    for i, (t, v, c, cond) in enumerate(rows_data):
        r = price_table.rows[i+1].cells
        r[0].text, r[1].text, r[2].text, r[3].text = t, v, c, cond

    # FOOTER
    doc.add_paragraph()
    pf = doc.add_paragraph('Proyecto academico basado en implementacion real - El Mariachi Aventurero v4.0')
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.runs[0].font.size = Pt(8)
    pf.runs[0].font.color.rgb = RGBColor(130, 130, 130)

    doc.save(OUTPUT)
    print(f'OK - Documento guardado: {OUTPUT}')

if __name__ == '__main__':
    main()
