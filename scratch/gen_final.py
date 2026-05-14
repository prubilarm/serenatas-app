from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

IMG_PATH = r'C:\Users\prubi\.gemini\antigravity\brain\7387ae71-aba8-4bf4-b378-d6579c5aeff8\mer_serenatas_profesional_1778780617780.png'
OUTPUT = 'Trabajo_Final_Serenatas_MER.docx'

def heading(doc, text, level):
    h = doc.add_heading(text, level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def justified(doc, *runs):
    """Add a justified paragraph with multiple runs: (text, bold, underline)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for item in runs:
        if isinstance(item, str):
            p.add_run(item)
        elif isinstance(item, dict):
            r = p.add_run(item['text'])
            r.bold = item.get('bold', False)
            r.underline = item.get('underline', False)
            r.italic = item.get('italic', False)
    return p

def bullet(doc, label, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f'{label}: ')
    r.bold = True
    p.add_run(text)
    return p

def main():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # === ENCABEZADO ===
    ph = doc.add_paragraph()
    ph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rh = ph.add_run('Actividad de Aula — Abstracción de Datos\nProyecto: El Mariachi Aventurero v4.0')
    rh.font.size = Pt(9)
    rh.font.color.rgb = RGBColor(100, 100, 100)

    # ==========================================
    # CARA A
    # ==========================================
    ta = doc.add_heading('A     CARA FRONTAL: HISTORIA Y PROCESOS', 0)
    ta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. CONTEXTO
    heading(doc, '1. Contexto del Proyecto', 1)
    justified(doc,
        'El proyecto ',
        {'text': '"El Mariachi Aventurero"', 'bold': True},
        ' es una solución tecnológica de gestión interna diseñada para el músico-empresario. '
        'A diferencia de plataformas de reserva para el público, esta herramienta está orientada exclusivamente '
        'al ',
        {'text': 'administrador', 'underline': True},
        ' del negocio, es decir, ',
        {'text': 'el propio músico', 'bold': True},
        '. Su objetivo es eliminar el caos de gestionar ',
        {'text': 'serenatas', 'underline': True},
        ' por WhatsApp o anotaciones físicas, centralizando en una sola plataforma digital los datos de sus ',
        {'text': 'clientes', 'underline': True},
        ', los detalles logísticos de cada presentación y el flujo de ',
        {'text': 'pagos', 'underline': True},
        ' con sus respectivos ',
        {'text': 'comprobantes', 'underline': True},
        '.'
    )

    justified(doc,
        'La plataforma cuenta con tres interfaces: una ',
        {'text': 'aplicación web (panel administrativo)', 'bold': True},
        ', una ',
        {'text': 'aplicación móvil', 'bold': True},
        ' para el músico en terreno, y una ',
        {'text': 'API REST', 'bold': True},
        ' que conecta ambas con la base de datos en Supabase. Así, el ',
        {'text': 'músico', 'underline': True},
        ' puede verificar su agenda desde el celular mientras se traslada entre ',
        {'text': 'serenatas', 'underline': True},
        ' y consultar reportes de ingresos desde su computador.'
    )

    # 2. PROCESOS
    heading(doc, '2. Procesos del Sistema (Narrativa del Flujo de Datos)', 1)

    justified(doc,
        'El sistema opera bajo un flujo secuencial que cubre el ciclo de vida completo de cada servicio musical. '
        'Cada vez que el ',
        {'text': 'músico', 'underline': True},
        ' recibe una solicitud, se desencadena una cadena de procesos que va desde el primer contacto hasta el cierre financiero:'
    )

    bullet(doc, 'Registro del Cliente',
        'El proceso inicia cuando se capta una solicitud. El sistema verifica si el contacto '
        'ya existe por teléfono; si no, crea un nuevo perfil de cliente con nombre, teléfono '
        'y observaciones opcionales.')

    bullet(doc, 'Agendamiento de la Serenata',
        'Se crea un registro de serenata vinculado al cliente. Se capturan todos los datos '
        'logísticos: nombre de la festejada, motivo, fecha, hora exacta, dirección, '
        'comuna, tipo de servicio (express o full) y precio total acordado.')

    bullet(doc, 'Seguimiento por Estados',
        'La serenata avanza por un ciclo de estados: consulta → cotizada → confirmada → '
        '"en camino" → realizada → pagada / cancelada. Esto permite al músico saber '
        'exactamente qué eventos requieren acción inmediata.')

    bullet(doc, 'Registro de Transacciones',
        'Al concretar el servicio, se registra el pago indicando monto, método '
        '(efectivo o transferencia) y se adjunta la URL del comprobante digital almacenado '
        'en Supabase Storage.')

    bullet(doc, 'Generación de Reportes',
        'El sistema consolida pagos y serenatas para mostrar ingresos por período, '
        'comunas con más demanda, y efectividad de cobros.')

    # 3. ENTIDADES
    heading(doc, '3. Identificación de Entidades (Sustantivos Clave)', 1)
    justified(doc,
        'Los siguientes sustantivos extraídos de la narrativa representan las entidades del modelo relacional:'
    )

    entities = [
        ('Cliente', 'Persona que contacta y solicita el servicio musical.'),
        ('Serenata', 'El evento o presentación musical central del negocio.'),
        ('Pago', 'Transacción económica que cierra el ciclo de servicio.'),
        ('Músico', 'El administrador del sistema (propietario del negocio).'),
        ('Festejada', 'Persona que recibe el homenaje (dato de la serenata).'),
        ('Comprobante', 'Documento digital que respalda un pago registrado.'),
        ('Estado', 'Condición lógica que define la etapa actual del servicio.'),
    ]
    for name, desc in entities:
        p = doc.add_paragraph()
        r_name = p.add_run(f'• {name}')
        r_name.bold = True
        r_name.underline = True
        p.add_run(f': {desc}')

    # ==========================================
    # SALTO DE PÁGINA — CARA B
    # ==========================================
    doc.add_page_break()

    tb = doc.add_heading('B     CARA POSTERIOR: DIAGRAMA MER', 0)
    tb.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading(doc, '1. Diagramación — Notación de Chen', 1)
    justified(doc,
        'El siguiente diagrama representa el Modelo Entidad-Relación (MER) completo del sistema, '
        'utilizando la notación de Chen: rectángulos para entidades, óvalos para atributos y '
        'rombos para las relaciones. Los atributos marcados como ',
        {'text': 'PK', 'bold': True},
        ' (subrayados) corresponden a las llaves primarias, y los marcados como ',
        {'text': 'FK', 'bold': True},
        ' son las llaves foráneas que vinculan las tablas.'
    )

    doc.add_paragraph()

    # Insertar imagen
    try:
        doc.add_picture(IMG_PATH, width=Inches(6.2))
        cap = doc.add_paragraph('Diagrama M.E.R. — Sistema El Mariachi Aventurero v4.0')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].italic = True
    except Exception as e:
        doc.add_paragraph(f'[ERROR insertando imagen: {e}]')

    doc.add_paragraph()

    # 2. ATRIBUTOS Y PK
    heading(doc, '2. Atributos y Llaves Primarias / Foráneas', 1)

    attr_data = [
        ('CLIENTE', [
            ('id', 'UUID', 'PK'),
            ('nombre', 'TEXT', ''),
            ('telefono', 'TEXT', ''),
            ('observaciones', 'TEXT', ''),
            ('created_at', 'TIMESTAMP', ''),
        ]),
        ('SERENATA', [
            ('id', 'UUID', 'PK'),
            ('cliente_id', 'UUID', 'FK → CLIENTE'),
            ('nombre_festejada', 'TEXT', ''),
            ('motivo', 'TEXT', ''),
            ('fecha', 'DATE', ''),
            ('hora', 'TIME', ''),
            ('direccion', 'TEXT', ''),
            ('comuna', 'TEXT', ''),
            ('mensaje_especial', 'TEXT', ''),
            ('tipo', "TEXT ('express'|'full')", ''),
            ('precio_total', 'NUMERIC', ''),
            ('estado', "TEXT ('consulta'...'pagada')", ''),
            ('created_at', 'TIMESTAMP', ''),
        ]),
        ('PAGO', [
            ('id', 'UUID', 'PK'),
            ('serenata_id', 'UUID', 'FK → SERENATA'),
            ('monto', 'NUMERIC', ''),
            ('metodo', "TEXT ('efectivo'|'transferencia')", ''),
            ('comprobante_url', 'TEXT', ''),
            ('fecha_pago', 'TIMESTAMP', ''),
        ]),
    ]

    for entity_name, fields in attr_data:
        doc.add_paragraph()
        ph_entity = doc.add_paragraph()
        ph_entity.add_run(f'Entidad: {entity_name}').bold = True

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = 'Atributo', 'Tipo de Dato', 'Llave'
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True

        for fname, ftype, fkey in fields:
            row = table.add_row().cells
            r0 = row[0].paragraphs[0].add_run(fname)
            if fkey.startswith('PK'):
                r0.underline = True
                r0.bold = True
            row[1].text = ftype
            row[2].text = fkey

    # 3. CARDINALIDADES
    doc.add_paragraph()
    heading(doc, '3. Cardinalidades y Reglas de Negocio', 1)

    cards = [
        ('CLIENTE — SERENATA  (1:N)',
         'Un mismo cliente puede agendar múltiples serenatas a lo largo del tiempo. '
         'Sin embargo, cada serenata registrada pertenece a un único cliente. '
         'La eliminación de un cliente deja la serenata sin vínculo (SET NULL).'),
        ('SERENATA — PAGO  (1:N)',
         'Una serenata puede recibir varios pagos, permitiendo gestionar abonos parciales '
         'y el pago del saldo restante. Si la serenata es eliminada, sus pagos asociados '
         'se eliminan en cascada (ON DELETE CASCADE).'),
    ]
    for title_c, desc_c in cards:
        p_c = doc.add_paragraph()
        p_c.add_run(f'• {title_c}: ').bold = True
        p_c.add_run(desc_c)

    # FOOTER
    doc.add_paragraph()
    pf = doc.add_paragraph('Proyecto académico basado en implementación real — El Mariachi Aventurero v4.0')
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.runs[0].font.size = Pt(8)
    pf.runs[0].font.color.rgb = RGBColor(130, 130, 130)

    doc.save(OUTPUT)
    print(f'OK - Documento guardado: {OUTPUT}')

if __name__ == '__main__':
    main()
